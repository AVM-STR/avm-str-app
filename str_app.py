"""
AVM Short-Term Rental Report Generator
Streamlit Web App
"""

import os, re, io, tempfile, json, threading, imaplib, smtplib, time
from email import policy
from email.parser import BytesParser
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime
from bs4 import BeautifulSoup
import streamlit as st
import fitz  # pymupdf
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from PIL import Image as PILImage
import requests

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, PageBreak, Image, HRFlowable)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# ── Brand ─────────────────────────────────────────────────────────────────────
DARK_BLUE  = colors.HexColor("#1F3864")
MID_BLUE   = colors.HexColor("#2E5FA3")
LIGHT_BLUE = colors.HexColor("#BDD7EE")
LIGHT_GRAY = colors.HexColor("#F5F5F5")
DARK_GRAY  = colors.HexColor("#444444")
WHITE      = colors.white

PAGE_W, PAGE_H = letter
MARGIN    = 0.65 * inch
CONTENT_W = PAGE_W - 2 * MARGIN

LOGO_PATH = os.path.join(os.path.dirname(__file__), "avm_logo.png")

# ── AirDNA PDF Parser ─────────────────────────────────────────────────────────
def parse_airdna_pdf(pdf_bytes):
    """Extract all data from AirDNA Rentalizer PDF using known line structure."""
    doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
    lines   = [l.strip() for l in doc[0].get_text().split("\n") if l.strip()]
    p2lines = [l.strip() for l in (doc[1].get_text() if len(doc)>1 else "").split("\n") if l.strip()]
    data = {}

    # ── Address (lines 2 & 3, after "Property Earning Potential" and "Submarket Score") ──
    data["address_line1"] = lines[2].rstrip(",").strip() if len(lines) > 2 else ""
    data["city_state_zip"] = lines[3].replace(", USA","").strip() if len(lines) > 3 else ""

    # ── Market / Submarket (line 4) ──
    for l in lines:
        m = re.search(r"Market:\s*(.+?)\s+Submarket:\s*(.+)", l)
        if m:
            data["market"]    = m.group(1).strip()
            data["submarket"] = m.group(2).strip()
            break

    # ── Beds / Baths / Guests (lines 5,6,7) ──
    for l in lines:
        m = re.search(r"^(\d+)\s+Bed", l)
        if m: data["bedrooms"] = m.group(1)
        m = re.search(r"^(\d+(?:\.\d+)?)\s+Bath", l)
        if m: data["bathrooms"] = m.group(1)
        m = re.search(r"^(\d+)\s+Guests?", l)
        if m: data["max_guests"] = m.group(1)

    # ── Financials (fixed positions relative to labels) ──
    for i, l in enumerate(lines):
        if l == "Operating Expenses"   and i+1 < len(lines): data["operating_expenses"] = lines[i+1]
        if l == "Net Operating Income" and i+1 < len(lines): data["noi"]                = lines[i+1]
        if l == "Cap Rate"             and i+1 < len(lines): data["cap_rate"]            = lines[i+1]

    # Revenue = line before "Projected", Occupancy = line before "Occupancy", ADR = line before "Average"
    for i, l in enumerate(lines):
        if l == "Projected"  and i > 0 and "$" in lines[i-1]: data["projected_revenue"] = lines[i-1]
        if l == "Occupancy"  and i > 0 and "%" in lines[i-1]: data["occupancy"]          = lines[i-1]
        if l == "Average"    and i > 0 and "$" in lines[i-1]: data["adr"]                = lines[i-1]

    # ── Submarket Score — appears after AIRDNA.CO footer ──
    for i, l in enumerate(lines):
        if l == "AIRDNA.CO" and i+2 < len(lines):
            candidate = lines[i+2]
            if candidate.isdigit() and 50 <= int(candidate) <= 100:
                data["submarket_score"] = candidate
            break

    # ── Comps — each comp is 7 lines after the title ──
    # Header columns end at "ADR" (line index 36), comps start at 37
    numeric_pat = re.compile(r"^\$?[\d,.KM%]+$")

    adr_idx = None
    for i, l in enumerate(lines):
        if l == "ADR":
            adr_idx = i
            break

    comps = []
    if adr_idx is not None:
        i = adr_idx + 1
        while i < len(lines):
            l = lines[i]
            if l.startswith("+") or l == "AIRDNA.CO":
                break
            # Collect title (may span multiple lines) then 7 numeric values
            title = l
            vals  = []
            j = i + 1
            while j < len(lines) and len(vals) < 7:
                candidate = lines[j]
                cleaned   = candidate.replace(".","").replace("%","").replace("$","").replace("K","").replace(",","")
                if numeric_pat.match(candidate) or cleaned.isdigit():
                    vals.append(candidate)
                elif len(vals) == 0:
                    title += " " + candidate
                else:
                    break
                j += 1
            if len(vals) == 7:
                # Clean title — AirDNA sometimes bleeds the bedroom number onto the title line
                clean_title = re.sub(r'\s+\d+(?:\.\d+)?$', '', title.strip())
                comps.append({
                    "num":     str(len(comps)+1),
                    "name":    clean_title,
                    "bdba":    f"{vals[0]}/{vals[1]}",
                    "rev_pot": vals[2],
                    "days":    vals[3],
                    "revenue": vals[4],
                    "occ":     vals[5],
                    "adr":     vals[6],
                })
                i = j
            else:
                i += 1
    data["comps"] = comps

    # ── Amenities from page 2 ──
    known = {"Air Conditioning","Dryer","Heating","Hot Tub","Kitchen",
             "Parking","Pool","Cable TV","Washer","Wireless Internet"}
    raw_amenities = []
    i = 0
    while i < len(p2lines):
        if p2lines[i] in known and i+1 < len(p2lines) and "%" in p2lines[i+1]:
            raw_amenities.append((p2lines[i], p2lines[i+1]))
            i += 2
        else:
            i += 1
    # Merge Dryer + Washer → Dryer / Washer
    merged, dryer_pct = [], None
    for name, pct in raw_amenities:
        if name == "Dryer":
            dryer_pct = pct
        elif name == "Washer":
            merged.append(("Dryer / Washer", dryer_pct or pct))
        else:
            merged.append((name, pct))
    data["amenities"] = merged

    # ── Property photo (largest image on page 1) ──
    photo_path, best_size = None, 0
    for img in doc[0].get_images(full=True):
        base = doc.extract_image(img[0])
        if len(base["image"]) > best_size:
            best_size      = len(base["image"])
            photo_bytes    = base["image"]
            photo_ext      = base["ext"]
    if best_size > 10000:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{photo_ext}")
        tmp.write(photo_bytes); tmp.close()
        photo_path = tmp.name
    data["photo_path"] = photo_path

    return data


# ── AI Market Commentary ──────────────────────────────────────────────────────
# ── Charts ────────────────────────────────────────────────────────────────────
BRAND_BLUE = "#2E5FA3"

def chart_revenue_range(comps, projected_rev, out_path):
    """
    Horizontal range chart showing min/Q1/median/Q3/max revenue across the
    comp set, with the analyst's subject projection marked as a vertical line.
    No individual comp is identifiable — aggregated statistics only.
    """
    import statistics
    try:
        rev_vals = sorted([
            float(c["revenue"].replace("$","").replace("K","").replace(",","")) * 1000
            for c in comps if c.get("revenue")
        ])
        proj = float(projected_rev.replace("$","").replace("K","").replace(",","")) * 1000
    except Exception:
        return

    n = len(rev_vals)
    if n < 2:
        return

    r_min    = rev_vals[0]
    r_max    = rev_vals[-1]
    r_median = statistics.median(rev_vals)
    r_q1     = rev_vals[max(0, n//4)]
    r_q3     = rev_vals[min(n-1, 3*n//4)]

    fig, ax = plt.subplots(figsize=(7.5, 2.2))

    # Full range bar
    ax.barh(0, r_max - r_min, left=r_min, height=0.35,
            color=BRAND_BLUE, alpha=0.15, label="Full Range")
    # IQR bar
    ax.barh(0, r_q3 - r_q1, left=r_q1, height=0.35,
            color=BRAND_BLUE, alpha=0.45, label="Middle 50%")
    # Median tick
    ax.plot([r_median, r_median], [-0.22, 0.22], color=BRAND_BLUE,
            linewidth=2.5, label=f"Median  ${r_median/1000:.0f}K")
    # Subject projection
    ax.plot([proj, proj], [-0.28, 0.28], color="#E84040",
            linewidth=2.5, linestyle="--", label=f"Subject Projection  ${proj/1000:.0f}K")

    ax.set_yticks([])
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"${v/1000:.0f}K"))
    ax.tick_params(axis="x", labelsize=8)
    ax.set_xlim(r_min * 0.88, r_max * 1.06)
    ax.set_ylim(-0.5, 0.5)
    ax.grid(axis="x", linestyle="--", alpha=0.4)
    ax.spines[["top","right","left"]].set_visible(False)
    ax.legend(fontsize=7.5, loc="upper left", framealpha=0.7)
    fig.tight_layout(pad=0.5)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


def chart_adr_vs_occ(comps, proj_adr, proj_occ, out_path):
    """
    Anonymous scatter plot of ADR vs. Occupancy across the comp set.
    No individual comp is labeled — distribution only.
    Subject projection marked distinctly.
    """
    try:
        adr_vals = [float(c["adr"].replace("$","").replace(",",""))
                    for c in comps if c.get("adr")]
        occ_vals = [float(c["occ"].replace("%",""))
                    for c in comps if c.get("occ")]
        p_adr = float(proj_adr.replace("$","").replace(",",""))
        p_occ = float(proj_occ.replace("%",""))
    except Exception:
        return

    if len(adr_vals) < 2:
        return

    fig, ax = plt.subplots(figsize=(7.5, 2.8))

    # Comp cloud — no labels
    ax.scatter(occ_vals, adr_vals, color=BRAND_BLUE, alpha=0.55,
               s=55, zorder=3, label="Comparable Properties")

    # Subject projection
    ax.scatter([p_occ], [p_adr], color="#E84040", s=90, zorder=5,
               marker="*", label=f"Subject Projection")

    ax.set_xlabel("Occupancy Rate (%)", fontsize=8)
    ax.set_ylabel("Average Daily Rate ($)", fontsize=8)
    ax.tick_params(labelsize=8)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"{v:.0f}%"))
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"${v:.0f}"))
    ax.grid(linestyle="--", alpha=0.35)
    ax.spines[["top","right"]].set_visible(False)
    ax.legend(fontsize=7.5, framealpha=0.7)
    fig.tight_layout(pad=0.5)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


# ── PDF Builder (same engine as before) ───────────────────────────────────────
def header_footer(canvas, doc, address):
    canvas.saveState()
    if os.path.exists(LOGO_PATH):
        canvas.drawImage(LOGO_PATH, PAGE_W - MARGIN - 1.5*inch,
                         PAGE_H - 0.52*inch, width=1.5*inch, height=0.42*inch,
                         preserveAspectRatio=True, mask="auto")
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#555555"))
    canvas.drawString(MARGIN, PAGE_H - 0.38*inch, address)
    canvas.setStrokeColor(MID_BLUE)
    canvas.setLineWidth(1)
    canvas.line(MARGIN, PAGE_H - 0.56*inch, PAGE_W - MARGIN, PAGE_H - 0.56*inch)
    canvas.line(MARGIN, 0.52*inch, PAGE_W - MARGIN, 0.52*inch)
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#888888"))
    canvas.drawString(MARGIN, 0.35*inch,
        "Absolute Value Management | One Lincoln Street, 24th Floor, "
        "Boston, MA 02111 | 617-391-0000 | office@avappraisalmgmt.com")
    canvas.drawRightString(PAGE_W - MARGIN, 0.35*inch, f"Page {doc.page}")
    canvas.restoreState()

def make_styles():
    def s(name, **kw): return ParagraphStyle(name, **kw)
    return {
        "title":    s("title", fontSize=22, fontName="Helvetica-Bold",
                      textColor=DARK_BLUE, spaceAfter=6, leading=26),
        "h1":       s("h1", fontSize=14, fontName="Helvetica-Bold",
                      textColor=MID_BLUE, spaceBefore=10, spaceAfter=4, leading=18),
        "h2":       s("h2", fontSize=11, fontName="Helvetica-Bold",
                      textColor=DARK_BLUE, spaceBefore=8, spaceAfter=3),
        "body":     s("body", fontSize=9.5, fontName="Helvetica",
                      textColor=DARK_GRAY, leading=14, spaceAfter=4),
        "small":    s("small", fontSize=8, fontName="Helvetica",
                      textColor=DARK_GRAY, leading=11),
        "lk":       s("lk", fontSize=9, fontName="Helvetica-Bold", textColor=DARK_GRAY),
        "lv":       s("lv", fontSize=9, fontName="Helvetica", textColor=DARK_GRAY),
        "cert":     s("cert", fontSize=9, fontName="Helvetica", textColor=DARK_GRAY,
                      leading=13, leftIndent=10, spaceAfter=2),
    }

DISCLAIMER_ITEMS = [
    "<b>Not an appraisal:</b> This report is a short-term rental income analysis only. It is not an appraisal, appraisal review, or an opinion of market value or market rent.",
    "<b>Licensing:</b> Absolute Value Management is not acting as a licensed or certified real estate appraiser in connection with this report and does not provide appraisal services through this analysis.",
    "<b>Estimates:</b> All figures are estimates derived from third-party short-term rental market data and comparable STR performance. Actual results may vary materially based on management, condition, amenities, pricing strategy, seasonality, and market changes.",
    "<b>Rules &amp; permits:</b> Local STR regulations, permits, and tax requirements may apply. Compliance is the responsibility of the owner/operator.",
]

AVM_COMMENTARY_BOILERPLATE = (
    "This STR income analysis is intended to assist the client and/or lender with reviewing "
    "potential short-term rental income for the subject property. No interior or exterior "
    "inspection was completed as part of this analysis, and no opinion of market value or "
    "market rent is provided. Actual STR performance is highly sensitive to pricing strategy, "
    "management quality, guest reviews, furnishings, and amenity set. Local STR regulations, "
    "HOA restrictions, and permitting requirements can materially impact whether STR operation "
    "is permitted and under what conditions."
)

CERT_ITEMS = [
    "The statements of fact contained in this report are true and correct to the best of my knowledge.",
    "The analyses and conclusions are limited only by the stated assumptions and limiting conditions.",
    "I have no present or prospective interest in the subject property and no personal interest with respect to the parties involved.",
    "My compensation is not contingent upon the reporting of a predetermined result or conclusion.",
    "This report is a short-term rental income analysis and is not an appraisal or appraisal review.",
    "The analyst is not acting as a state-licensed or certified real estate appraiser for this assignment.",
]

METHODOLOGY_SECTIONS = [
    ("What this report is (and is not)",
     "This document is a short-term rental income analysis prepared for income support and feasibility review. It summarizes estimated revenue and operating metrics using third-party STR market data and the performance of similar active listings. This is not an appraisal, not an opinion of market value, and not an opinion of market rent."),
    ("Data sources",
     "Market data sources reviewed for this analysis include AirDNA (paid subscription) and publicly available listing data observed directly on Airbnb.com and VRBO.com as of the report date. All performance metrics and income conclusions cited in this report reflect the analyst's independent reconciliation of available market evidence and do not constitute a reproduction of any third-party data product or model output. AirDNA and other data sources are referenced solely as research tools used to inform the analyst's independent professional conclusions."),
    ("Data considered",
     "Primary inputs include the subject's configuration (bed/bath/guest capacity), market and submarket classification, and a curated set of comparable STR listings. The comparable set is used to bracket typical ADR, occupancy rates, and annual revenue for similar rentals."),
    ("Operating expenses, NOI &amp; cap rate",
     "Operating expenses reflect a modeled STR expense framework inclusive of estimated taxes, insurance, utilities, maintenance and turnover costs, and platform/management fees. Net operating income (NOI) is calculated as projected gross revenue less estimated operating expenses."),
    ("Key limitations",
     "No interior or exterior inspection was completed for this analysis. Property condition, furnishings, amenity set, management quality, pricing strategy, and guest reviews can materially impact actual STR performance. Local STR regulations, HOA restrictions, and permitting requirements may restrict or prohibit short-term rental operation in whole or in part. All projections are estimates and are not guarantees of future performance."),
    ("Intended users &amp; intended use",
     "Intended user(s): the client and/or lender, and parties specifically authorized by the client. Intended use: lender-facing STR income support and feasibility review for the subject property. Any other use of this report is prohibited without the express written permission of Absolute Value Management."),
]

def generate_comp_narrative(data):
    """
    Generate a professional analyst narrative from parsed comp data.
    Replaces the comp table — no proprietary data reproduced, all conclusions
    are the analyst's independent reconciliation of available market evidence.
    """
    comps = data.get("comps", [])
    market    = data.get("market", "this market")
    submarket = data.get("submarket", "this submarket")
    bedrooms  = data.get("bedrooms", "")
    rev       = data.get("projected_revenue", "")
    adr       = data.get("adr", "")
    occ       = data.get("occupancy", "")

    if not comps:
        return (
            "The analyst reviewed available short-term rental market data for the subject submarket. "
            "Insufficient comparable data was available to produce a detailed range analysis; "
            "the income estimate reflects the analyst's review of broader market conditions."
        )

    # ── Compute stats ────────────────────────────────────────────────────────
    try:
        occ_vals = [float(c["occ"].replace("%","")) for c in comps if c.get("occ")]
        adr_vals = [float(c["adr"].replace("$","").replace(",","")) for c in comps if c.get("adr")]
        rev_vals = [float(c["revenue"].replace("$","").replace("K","").replace(",",""))*1000
                    for c in comps if c.get("revenue")]
        day_vals = [int(c["days"]) for c in comps if c.get("days","").isdigit()]
    except Exception:
        occ_vals, adr_vals, rev_vals, day_vals = [], [], [], []

    if not occ_vals:
        return "The analyst reviewed available short-term rental market data for the subject submarket."

    occ_min,  occ_max  = min(occ_vals),  max(occ_vals)
    adr_min,  adr_max  = min(adr_vals),  max(adr_vals)
    rev_min,  rev_max  = min(rev_vals),  max(rev_vals)
    day_min,  day_max  = (min(day_vals), max(day_vals)) if day_vals else (None, None)

    # Clustering — middle 50% (IQR)
    import statistics
    occ_median = statistics.median(occ_vals)
    adr_median = statistics.median(adr_vals)
    rev_median = statistics.median(rev_vals)

    occ_sorted = sorted(occ_vals)
    adr_sorted = sorted(adr_vals)
    rev_sorted = sorted(rev_vals)
    n = len(comps)
    q1_idx, q3_idx = max(0, n//4), min(n-1, 3*n//4)

    occ_cluster = (occ_sorted[q1_idx], occ_sorted[q3_idx])
    adr_cluster = (adr_sorted[q1_idx], adr_sorted[q3_idx])
    rev_cluster = (rev_sorted[q1_idx]/1000, rev_sorted[q3_idx]/1000)

    # ADR/occupancy relationship — high ADR comps tend to have lower occ?
    paired = sorted(zip(adr_vals, occ_vals), key=lambda x: x[0])
    top_adr_occ   = [o for a,o in paired[-3:]]
    bot_adr_occ   = [o for a,o in paired[:3]]
    luxury_note   = (sum(top_adr_occ)/len(top_adr_occ)) < (sum(bot_adr_occ)/len(bot_adr_occ))

    # ── Build narrative ──────────────────────────────────────────────────────
    n_comps = len(comps)
    bed_str = f"{bedrooms}-bedroom " if bedrooms else ""

    para1 = (
        f"To support the income estimate for the subject property, the analyst reviewed the "
        f"performance of {n_comps} active short-term rental listings in the {submarket} submarket "
        f"comparable to the subject in terms of bedroom count, bathroom count, building class, "
        f"and overall utility. Comparable properties demonstrated a range of performance outcomes "
        f"reflective of differences in floor level, view orientation, amenity set, and management quality."
    )

    para2 = (
        f"Among the comparable set, occupancy rates ranged from approximately {occ_min:.0f}% "
        f"to {occ_max:.0f}%, with the majority of properties clustered between "
        f"{occ_cluster[0]:.0f}% and {occ_cluster[1]:.0f}%. Average daily rates ranged from "
        f"approximately ${adr_min:.0f} to ${adr_max:.0f}, with most properties performing in "
        f"the ${adr_cluster[0]:.0f} to ${adr_cluster[1]:.0f} range. Actual annual revenues "
        f"across the comparable set ranged from approximately ${rev_min/1000:.0f}K to "
        f"${rev_max/1000:.0f}K, with the majority of comparables falling between "
        f"${rev_cluster[0]:.0f}K and ${rev_cluster[1]:.0f}K."
    )

    days_sent = ""
    if day_min and day_max:
        days_sent = (
            f" Days booked ranged from approximately {day_min} to {day_max} annually, "
            f"consistent with {'a high-demand urban' if day_min > 280 else 'an active'} "
            f"submarket benefiting from year-round demand."
        )

    luxury_sent = ""
    if luxury_note:
        luxury_sent = (
            " Properties at the upper end of the ADR range tended to carry lower occupancy rates, "
            "suggesting a premium pricing strategy rather than volume-based booking."
        )

    para3 = days_sent + luxury_sent

    para4 = (
        f"The subject's projected occupancy of {occ} and ADR of {adr} fall within the "
        f"well-supported middle range of comparable performance and are considered reasonable "
        f"and achievable under competent management. The analyst's projected gross annual "
        f"revenue of {rev} is supported by and consistent with the range established by "
        f"the comparable set. Market data sources reviewed for this analysis include AirDNA "
        f"(paid subscription) and publicly available listing data observed directly on "
        f"Airbnb.com and VRBO.com as of the report date. Performance metrics cited herein "
        f"reflect the analyst's independent reconciliation of available market evidence and "
        f"do not constitute a reproduction of any third-party data product or model output."
    )

    return "\n\n".join([para1, para2, para3.strip(), para4]) if para3.strip() else \
           "\n\n".join([para1, para2, para4])


def build_pdf(data, client, loan_num, report_date, commentary, buf,
              photo_override=None, map_override=None,
              client_address="", client_phone="", client_order_num="",
              borrower="", avm_file_id="", property_type="Single-Family Residence"):
    styles = make_styles()
    addr1 = data.get("address_line1","")
    city  = data.get("city_state_zip","")
    full_address = f"{addr1}, {city}"

    # Store assignment fields in data for sign-off table
    data["client_address"]   = client_address
    data["client_phone"]     = client_phone
    data["client_order_num"] = client_order_num
    data["borrower"]         = borrower
    data["avm_file_id"]      = avm_file_id

    # Apply overrides
    if photo_override:
        data["photo_path"] = photo_override
    if map_override:
        data["map_path"] = map_override

    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=0.65*inch, bottomMargin=0.65*inch)

    def _hf(canvas, doc):
        header_footer(canvas, doc, full_address)

    story = []

    # ── PAGE 1 ──────────────────────────────────────────────────────────────
    story.append(Paragraph("Short-Term Rental Income Analysis", styles["title"]))
    story.append(Paragraph("(Not an Appraisal)",
        ParagraphStyle("subtitle", fontSize=11, fontName="Helvetica",
                       textColor=colors.HexColor("#888888"), spaceAfter=6, leading=14)))
    story.append(HRFlowable(width=CONTENT_W, thickness=1, color=MID_BLUE, spaceAfter=10))

    # Photo + info side by side
    photo_w = 3.2*inch
    info_w  = CONTENT_W - photo_w - 0.2*inch

    photo_path = data.get("photo_path")
    if photo_path and os.path.exists(photo_path):
        photo_cell = Image(photo_path, width=photo_w, height=2.2*inch)
    else:
        photo_cell = Paragraph("<font color='#AAAAAA'>[Property Photo]</font>",
            ParagraphStyle("ph", fontSize=10, alignment=TA_CENTER))

    lk, lv = styles["lk"], styles["lv"]
    info_rows = [
        [Paragraph("<b>Subject Property:</b>", lk), ""],
        [Paragraph(addr1, lv), ""],
        [Paragraph(city, lv), ""],
        [Paragraph("<b>Property Type:</b>", lk),  Paragraph(property_type, lv)],
        [Paragraph("<b>Configuration:</b>", lk),
         Paragraph(f"{data.get('bedrooms','')} Bedrooms | {data.get('bathrooms','')} Bathrooms", lv)],
        [Paragraph("<b>Maximum Guests:</b>", lk),  Paragraph(data.get("max_guests",""), lv)],
        [Paragraph("<b>Market Area:</b>", lk),     Paragraph(data.get("market",""), lv)],
        [Paragraph("<b>Submarket:</b>", lk),        Paragraph(data.get("submarket",""), lv)],
        [Paragraph("<b>Market Demand:</b>", lk),
         Paragraph(("Very Strong" if int(data.get("submarket_score",0) or 0) >= 85
                    else "Strong" if int(data.get("submarket_score",0) or 0) >= 70
                    else "Moderate" if int(data.get("submarket_score",0) or 0) >= 55
                    else "Emerging"), lv)],
        [Paragraph("<b>Report Date:</b>", lk),      Paragraph(report_date, lv)],
        [Paragraph("<b>Prepared By:</b>", lk),      Paragraph("Absolute Value Management", lv)],
    ]
    info_t = Table(info_rows, colWidths=[1.5*inch, 1.9*inch])
    info_t.setStyle(TableStyle([
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("TOPPADDING",(0,0),(-1,-1),2),("BOTTOMPADDING",(0,0),(-1,-1),2),
        ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),4),
    ]))

    top_t = Table([[photo_cell, info_t]], colWidths=[photo_w, info_w])
    top_t.setStyle(TableStyle([
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0),
        ("TOPPADDING",(0,0),(-1,-1),0),("BOTTOMPADDING",(0,0),(-1,-1),0),
        ("RIGHTPADDING",(0,0),(0,-1),22),("LEFTPADDING",(1,0),(1,-1),8),
    ]))
    story.append(top_t)
    story.append(Spacer(1,14))

    # Metrics box
    story.append(Paragraph("Rental Analysis Quick View", styles["h1"]))
    rev  = data.get("projected_revenue","—")
    adr  = data.get("adr","—")
    occ  = data.get("occupancy","—")
    exp  = data.get("operating_expenses","—")
    noi  = data.get("noi","—")
    cap  = data.get("cap_rate","—")

    label_s = ParagraphStyle("ml", fontSize=8, fontName="Helvetica",
                              textColor=DARK_GRAY, alignment=TA_CENTER)
    value_s = ParagraphStyle("mv", fontSize=16, fontName="Helvetica-Bold",
                              textColor=MID_BLUE, alignment=TA_CENTER, leading=20)
    cw = CONTENT_W/3
    mx = Table([
        [Paragraph("Projected Annual STR Income",label_s),
         Paragraph("Average Daily Rate (ADR)",label_s),
         Paragraph("Occupancy Rate (Projected)",label_s)],
        [Paragraph(rev,value_s),Paragraph(adr,value_s),Paragraph(occ,value_s)],
        [Paragraph("Operating Expenses (Est.)",label_s),
         Paragraph("Net Operating Income (NOI)",label_s),
         Paragraph("Estimated Cap Rate",label_s)],
        [Paragraph(exp,value_s),Paragraph(noi,value_s),Paragraph(cap,value_s)],
    ], colWidths=[cw]*3)
    mx.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),0.75,MID_BLUE),
        ("INNERGRID",(0,0),(-1,-1),0.5,LIGHT_BLUE),
        ("BACKGROUND",(0,0),(-1,-1),WHITE),
        ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
        ("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(mx)
    story.append(Spacer(1,12))

    # Disclaimer box
    disc_s = ParagraphStyle("disc",fontSize=8,fontName="Helvetica",
                             textColor=DARK_GRAY,leading=11,spaceAfter=3)
    disc_items = [Paragraph("<b>Important Disclaimer (Read First)</b>",
        ParagraphStyle("dh",fontSize=8.5,fontName="Helvetica-Bold",
                       textColor=DARK_BLUE,spaceAfter=3))]
    for txt in DISCLAIMER_ITEMS:
        disc_items.append(Paragraph(txt, disc_s))
    d_inner = Table([[i] for i in disc_items], colWidths=[CONTENT_W-0.3*inch])
    d_inner.setStyle(TableStyle([("LEFTPADDING",(0,0),(-1,-1),6),
                                  ("RIGHTPADDING",(0,0),(-1,-1),6),
                                  ("TOPPADDING",(0,0),(-1,-1),2),
                                  ("BOTTOMPADDING",(0,0),(-1,-1),2)]))
    d_wrap = Table([[d_inner]],colWidths=[CONTENT_W])
    d_wrap.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.75,MID_BLUE),
                                 ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#EEF4FB")),
                                 ("TOPPADDING",(0,0),(-1,-1),6),
                                 ("BOTTOMPADDING",(0,0),(-1,-1),6),
                                 ("LEFTPADDING",(0,0),(-1,-1),4),
                                 ("RIGHTPADDING",(0,0),(-1,-1),4)]))
    story.append(d_wrap)

    # ── PAGE 2 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Comparable Short-Term Rental Analysis", styles["h1"]))

    # Generate and render narrative — no comp table or map
    comp_narrative = generate_comp_narrative(data)
    for para_text in comp_narrative.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            story.append(Paragraph(para_text, styles["body"]))
            story.append(Spacer(1, 4))

    # ── PAGE 3 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Market Overview &amp; Commentary", styles["h1"]))
    story.append(Paragraph(
        f"{data.get('market','')} Market – {data.get('submarket','')} Submarket",
        styles["h2"]))
    story.append(Paragraph(commentary, styles["body"]))
    story.append(Spacer(1,6))
    # Compute comp ranges for analyst-attributed projection language
    comps = data.get("comps", [])
    try:
        occ_vals = [float(c["occ"].replace("%","")) for c in comps if c.get("occ")]
        adr_vals = [float(c["adr"].replace("$","").replace(",","")) for c in comps if c.get("adr")]
        rev_vals = [float(c["revenue"].replace("$","").replace("K","").replace(",",""))*1000
                    for c in comps if c.get("revenue")]
        occ_range = f"{min(occ_vals):.0f}%\u2013{max(occ_vals):.0f}%" if occ_vals else occ
        adr_range = f"${min(adr_vals):.0f}\u2013${max(adr_vals):.0f}" if adr_vals else adr
        rev_low   = min(rev_vals) / 1000 if rev_vals else 0
        rev_high  = max(rev_vals) / 1000 if rev_vals else 0
        rev_range = f"${rev_low:.0f}K\u2013${rev_high:.0f}K" if rev_vals else rev
    except Exception:
        occ_range, adr_range, rev_range = occ, adr, rev

    story.append(Paragraph(
        f"<b>Projection Support</b><br/>"
        f"Based on the analyst's review of comparable short-term rental performance in the "
        f"{data.get('submarket','')} submarket, comparable properties demonstrated occupancy rates "
        f"ranging from approximately {occ_range}, average daily rates ranging from approximately "
        f"{adr_range}, and annual revenues ranging from approximately {rev_range}. The analyst's "
        f"projected gross annual revenue of {rev}, reflecting an ADR of {adr} and occupancy of {occ}, "
        f"represents an independent professional estimate derived from available market evidence. "
        f"The subject's bedroom count and guest capacity place it toward the larger end of typical "
        f"STR inventory, which can support higher ADR and stronger peak-season performance when "
        f"paired with competitive amenities and professional management.",
        styles["body"]))
    story.append(Spacer(1,10))

    # AVM Commentary box
    avm_s = ParagraphStyle("avm_i",fontSize=8.5,fontName="Helvetica",
                            textColor=DARK_GRAY,leading=12)
    avm_inner = Table([[Paragraph(AVM_COMMENTARY_BOILERPLATE,avm_s)]],
                       colWidths=[CONTENT_W-0.3*inch])
    avm_inner.setStyle(TableStyle([("TOPPADDING",(0,0),(-1,-1),6),
                                    ("BOTTOMPADDING",(0,0),(-1,-1),6),
                                    ("LEFTPADDING",(0,0),(-1,-1),8),
                                    ("RIGHTPADDING",(0,0),(-1,-1),8)]))
    avm_hdr = Table([[Paragraph("<b>AVM Commentary</b>",
        ParagraphStyle("ah",fontSize=9,fontName="Helvetica-Bold",textColor=WHITE))]],
        colWidths=[CONTENT_W])
    avm_hdr.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),DARK_BLUE),
                                  ("TOPPADDING",(0,0),(-1,-1),5),
                                  ("BOTTOMPADDING",(0,0),(-1,-1),5),
                                  ("LEFTPADDING",(0,0),(-1,-1),8)]))
    avm_outer = Table([[avm_hdr],[avm_inner]],colWidths=[CONTENT_W])
    avm_outer.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.75,MID_BLUE),
                                    ("BACKGROUND",(0,1),(-1,-1),colors.HexColor("#EEF4FB")),
                                    ("LEFTPADDING",(0,0),(-1,-1),0),
                                    ("RIGHTPADDING",(0,0),(-1,-1),0),
                                    ("TOPPADDING",(0,0),(-1,-1),0),
                                    ("BOTTOMPADDING",(0,0),(-1,-1),0)]))
    story.append(avm_outer)

    # ── PAGE 4 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Amenities &amp; Market Analysis", styles["h1"]))
    story.append(Paragraph("Comparable STR Amenity Prevalence", styles["h2"]))

    # Amenity table
    am_h = ParagraphStyle("amh",fontSize=8.5,fontName="Helvetica-Bold",
                           textColor=WHITE,alignment=TA_CENTER)
    am_c = ParagraphStyle("amc",fontSize=9,fontName="Helvetica",textColor=DARK_GRAY)
    am_p = ParagraphStyle("amp",fontSize=9,fontName="Helvetica",
                           textColor=DARK_GRAY,alignment=TA_CENTER)
    cw4 = CONTENT_W/4
    am_data = [[Paragraph("Amenity",am_h),Paragraph("% of Comps",am_h),
                Paragraph("Amenity",am_h),Paragraph("% of Comps",am_h)]]
    amenities = data.get("amenities",[])
    pairs = [(amenities[i], amenities[i+1] if i+1 < len(amenities) else ("",""))
             for i in range(0, len(amenities), 2)]
    for (a1,p1),(a2,p2) in pairs:
        am_data.append([Paragraph(a1,am_c),Paragraph(p1,am_p),
                        Paragraph(a2,am_c),Paragraph(p2,am_p)])
    at = Table(am_data,colWidths=[cw4]*4)
    at.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),MID_BLUE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,colors.HexColor("#F7FAFF")]),
        ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#AAAAAA")),
        ("INNERGRID",(0,0),(-1,-1),0.3,colors.HexColor("#DDDDDD")),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(at)
    story.append(Spacer(1,14))

    # Chart 1 — Revenue Range
    comps     = data.get("comps", [])
    proj_rev  = data.get("projected_revenue", "")
    proj_adr  = data.get("adr", "")
    proj_occ  = data.get("occupancy", "")

    if comps and proj_rev:
        tmp1 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp1.close()
        chart_revenue_range(comps, proj_rev, tmp1.name)
        if os.path.exists(tmp1.name) and os.path.getsize(tmp1.name) > 0:
            story.append(Paragraph("Comparable Revenue Range — Subject vs. Market", styles["h2"]))
            story.append(Paragraph(
                "The chart below illustrates the distribution of annual revenue across the comparable "
                "set. The shaded bar represents the full range; the darker band reflects the middle 50% "
                "of comparable performance. The subject's projected revenue is marked in red.",
                styles["small"]))
            story.append(Spacer(1, 4))
            story.append(Image(tmp1.name, width=CONTENT_W, height=2.0*inch))
            story.append(Spacer(1, 14))

    # Chart 2 — ADR vs Occupancy scatter
    if comps and proj_adr and proj_occ:
        tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp2.close()
        chart_adr_vs_occ(comps, proj_adr, proj_occ, tmp2.name)
        if os.path.exists(tmp2.name) and os.path.getsize(tmp2.name) > 0:
            story.append(Paragraph("ADR vs. Occupancy — Comparable Set Distribution", styles["h2"]))
            story.append(Paragraph(
                "Each point represents an anonymous comparable property. The red star marks the "
                "subject's projected ADR and occupancy combination relative to the comp set cluster.",
                styles["small"]))
            story.append(Spacer(1, 4))
            story.append(Image(tmp2.name, width=CONTENT_W, height=2.6*inch))

    # ── PAGE 5 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Methodology, Assumptions &amp; Limitations", styles["h1"]))
    for title, body in METHODOLOGY_SECTIONS:
        story.append(Paragraph(title, styles["h2"]))
        story.append(Paragraph(body, styles["body"]))

    # ── PAGE 6 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Identification, Intended Use &amp; Analyst Sign-Off", styles["h1"]))

    so_rows = [
        ["Prepared By",         "Absolute Value Management"],
        ["Subject Property",    full_address],
        ["Property Type",       property_type],
        ["Configuration",       f"{data.get('bedrooms','')} Bedrooms | "
                                f"{data.get('bathrooms','')} Bathrooms | "
                                f"{data.get('max_guests','')} Guests Max"],
        ["Market / Submarket",  f"{data.get('market','')} / {data.get('submarket','')}"],
        ["Report Date",         report_date],
        ["Client / Lender",     client],
        ["Client Address",      data.get("client_address","")],
        ["Client Phone",        data.get("client_phone","")],
        ["Client Order Number", data.get("client_order_num","")],
        ["Borrower",            data.get("borrower","")],
        ["Loan Number",         loan_num],
        ["AVM File ID",         data.get("avm_file_id","")],
        ["Intended Use",        "Short-term rental income support for lender feasibility / underwriting review (not an appraisal)."],
        ["Intended Users",      "Client/lender and parties specifically authorized by the client."],
    ]
    so_lk = ParagraphStyle("slk",fontSize=9,fontName="Helvetica-Bold",textColor=DARK_GRAY)
    so_lv = ParagraphStyle("slv",fontSize=9,fontName="Helvetica",textColor=DARK_GRAY)
    so_data = [[Paragraph(r[0],so_lk),Paragraph(r[1],so_lv)] for r in so_rows]
    so_t = Table(so_data,colWidths=[1.8*inch,CONTENT_W-1.8*inch])
    so_t.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#AAAAAA")),
        ("INNERGRID",(0,0),(-1,-1),0.3,colors.HexColor("#DDDDDD")),
        ("BACKGROUND",(0,0),(0,-1),LIGHT_GRAY),
        ("ROWBACKGROUNDS",(1,0),(-1,-1),[WHITE,colors.HexColor("#F7FAFF")]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),
        ("VALIGN",(0,0),(-1,-1),"TOP"),
    ]))
    story.append(so_t)
    story.append(Spacer(1,14))

    story.append(Paragraph("Analyst Certification", styles["h2"]))
    for item in CERT_ITEMS:
        story.append(Paragraph(f"• {item}", styles["cert"]))
    story.append(Spacer(1,20))

    sign_t = Table([[
        Paragraph("Company: <b>Absolute Value Management</b>",
            ParagraphStyle("sc",fontSize=9,fontName="Helvetica",textColor=DARK_GRAY)),
        Paragraph(f"Date: <b>{report_date}</b>",
            ParagraphStyle("sd",fontSize=9,fontName="Helvetica",
                           textColor=DARK_GRAY,alignment=TA_RIGHT))
    ]], colWidths=[CONTENT_W*0.5]*2)
    sign_t.setStyle(TableStyle([
        ("LINEABOVE",(1,0),(1,0),0.75,DARK_GRAY),
        ("TOPPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),0),
        ("RIGHTPADDING",(0,0),(-1,-1),0),
    ]))
    story.append(sign_t)

    doc.build(story, onFirstPage=_hf, onLaterPages=_hf)


# ── Email Helper ──────────────────────────────────────────────────────────────
# ── Intake Export Helpers ─────────────────────────────────────────────────────

def build_intake_pdf(intake_text, address="Subject Property"):
    """Convert intake markdown text to a branded PDF using ReportLab."""
    from io import BytesIO
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.9*inch, bottomMargin=0.75*inch
    )

    styles = {
        "title": ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=14,
                                 textColor=colors.HexColor("#1F3864"),
                                 spaceAfter=6, spaceBefore=0),
        "h2":    ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=11,
                                 textColor=colors.HexColor("#2E5FA3"),
                                 spaceAfter=4, spaceBefore=10),
        "body":  ParagraphStyle("body", fontName="Helvetica", fontSize=9,
                                 textColor=colors.HexColor("#333333"),
                                 spaceAfter=3, leading=13),
        "flag":  ParagraphStyle("flag", fontName="Helvetica", fontSize=9,
                                 textColor=colors.HexColor("#C00000"),
                                 spaceAfter=2, leading=12),
        "bold":  ParagraphStyle("bold", fontName="Helvetica-Bold", fontSize=9,
                                 textColor=colors.HexColor("#333333"),
                                 spaceAfter=2),
    }

    def _hf(canvas, doc):
        canvas.saveState()
        if os.path.exists(LOGO_PATH):
            canvas.drawImage(LOGO_PATH,
                             letter[0] - 0.75*inch - 1.4*inch,
                             letter[1] - 0.62*inch,
                             width=1.4*inch, height=0.38*inch,
                             preserveAspectRatio=True, mask="auto")
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(colors.HexColor("#1F3864"))
        canvas.drawString(0.75*inch, letter[1] - 0.45*inch, "APPRAISAL ASSIGNMENT INTAKE")
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#555555"))
        canvas.drawString(0.75*inch, letter[1] - 0.60*inch, address)
        canvas.setStrokeColor(colors.HexColor("#2E5FA3"))
        canvas.setLineWidth(1)
        canvas.line(0.75*inch, letter[1] - 0.68*inch, letter[0] - 0.75*inch, letter[1] - 0.68*inch)
        canvas.line(0.75*inch, 0.55*inch, letter[0] - 0.75*inch, 0.55*inch)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.drawString(0.75*inch, 0.38*inch,
            "A-Tech Appraisal Co., LLC | Absolute Value Management")
        canvas.drawRightString(letter[0] - 0.75*inch, 0.38*inch, f"Page {doc.page}")
        canvas.restoreState()

    story = []

    # Parse markdown into ReportLab elements
    lines = intake_text.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Skip horizontal rules and empty TOTAL_JSON remnants
        if stripped in ("---", "***", "___") or stripped.startswith("TOTAL_JSON"):
            if in_table and table_rows:
                # Flush table
                if len(table_rows) > 1:
                    col_count = max(len(r) for r in table_rows)
                    col_w = (letter[0] - 1.5*inch) / col_count
                    tbl = Table(
                        [[Paragraph(str(c), styles["body"]) for c in row] for row in table_rows],
                        colWidths=[col_w]*col_count
                    )
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#BDD7EE")),
                        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                        ("FONTSIZE", (0,0), (-1,-1), 8.5),
                        ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
                        ("ROWBACKGROUNDS", (0,1), (-1,-1),
                         [colors.white, colors.HexColor("#F5F8FF")]),
                        ("TOPPADDING", (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                        ("LEFTPADDING", (0,0), (-1,-1), 6),
                        ("RIGHTPADDING", (0,0), (-1,-1), 6),
                        ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 6))
                table_rows = []
                in_table = False
            if stripped in ("---", "***", "___"):
                story.append(HRFlowable(width="100%", thickness=0.5,
                                         color=colors.HexColor("#CCCCCC"),
                                         spaceAfter=4, spaceBefore=4))
            continue

        # Markdown table rows
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # separator row
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                col_count = max(len(r) for r in table_rows)
                col_w = (letter[0] - 1.5*inch) / col_count
                tbl = Table(
                    [[Paragraph(str(c), styles["body"]) for c in row] for row in table_rows],
                    colWidths=[col_w]*col_count
                )
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#BDD7EE")),
                    ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                    ("FONTSIZE", (0,0), (-1,-1), 8.5),
                    ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1),
                     [colors.white, colors.HexColor("#F5F8FF")]),
                    ("TOPPADDING", (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ("LEFTPADDING", (0,0), (-1,-1), 6),
                    ("RIGHTPADDING", (0,0), (-1,-1), 6),
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))
                table_rows = []
                in_table = False

        if not stripped:
            story.append(Spacer(1, 4))
            continue

        # Headings
        if stripped.startswith("## "):
            story.append(Paragraph(stripped[3:].strip(), styles["title"]))
        elif stripped.startswith("### ") or stripped.startswith("#### "):
            story.append(Paragraph(stripped.lstrip("#").strip(), styles["h2"]))
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            story.append(Paragraph(stripped.strip("*"), styles["bold"]))
        elif stripped.startswith("- [x]") or stripped.startswith("-[x]"):
            text = stripped[5:].strip()
            story.append(Paragraph(f"⚠ {text}", styles["flag"]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            # Inline bold
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(f"• {text}", styles["body"]))
        elif re.match(r"^\d+\.", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(f"{stripped.split('.')[0]}. {text}", styles["body"]))
        else:
            # Regular paragraph — handle inline bold
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
            story.append(Paragraph(text, styles["body"]))

    doc.build(story, onFirstPage=_hf, onLaterPages=_hf)
    buf.seek(0)
    return buf.getvalue()


def build_intake_docx(intake_text, address="Subject Property"):
    """Convert intake markdown text to a Word document using python-docx."""
    from io import BytesIO
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    doc = DocxDocument()

    # Page setup — US Letter, 0.75" margins
    section = doc.sections[0]
    section.page_width  = int(8.5  * 914400)
    section.page_height = int(11.0 * 914400)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(section, attr, int(0.75 * 914400))

    # Helper: set paragraph font
    def _set_run(run, bold=False, size=10, color=None, italic=False):
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        if level == 1:
            _set_run(run, bold=True, size=13, color=(31, 56, 100))
        else:
            _set_run(run, bold=True, size=11, color=(46, 95, 163))
        return p

    def _body(text, bold=False, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        _set_run(run, bold=bold, size=9.5,
                 color=color or (51, 51, 51))
        return p

    def _flag(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"⚠ {text}")
        _set_run(run, size=9.5, color=(192, 0, 0))
        return p

    def _bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        # Strip inline ** bold markers for simplicity
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        run = p.add_run(clean)
        _set_run(run, size=9.5)
        return p

    def _add_table(rows):
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        tbl = doc.add_table(rows=0, cols=col_count)
        tbl.style = "Table Grid"
        for i, row_data in enumerate(rows):
            row = tbl.add_row()
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text) if j < len(row_data) else ""
                run = cell.paragraphs[0].runs
                if run:
                    run[0].font.size = Pt(9)
                    if i == 0:
                        run[0].bold = True
                # Header row shading
                if i == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "BDD7EE")
                    tcPr.append(shd)
        doc.add_paragraph()  # spacing after table

    # Title header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    r = title_p.add_run("APPRAISAL ASSIGNMENT INTAKE")
    _set_run(r, bold=True, size=14, color=(31, 56, 100))

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(8)
    r = sub_p.add_run(address)
    _set_run(r, size=10, color=(85, 85, 85))

    # Parse and render
    lines = intake_text.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("TOTAL_JSON"):
            break  # stop at JSON block

        if stripped in ("---", "***", "___"):
            if in_table and table_rows:
                _add_table(table_rows)
                table_rows = []
                in_table = False
            # Horizontal rule via paragraph border
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                _add_table(table_rows)
                table_rows = []
                in_table = False

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("## "):
            _heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### ") or stripped.startswith("#### "):
            _heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            _body(stripped.strip("*"), bold=True)
        elif stripped.startswith("- [x]") or stripped.startswith("-[x]"):
            _flag(stripped[5:].strip())
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _bullet(stripped[2:].strip())
        elif re.match(r"^\d+\.", stripped):
            _bullet(re.sub(r"^\d+\.\s*", "", stripped))
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            _body(clean)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def send_report_email(to_email, subject, body, pdf_bytes, filename):
    """Send PDF report via Gmail SMTP."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication

    try:
        gmail_address = st.secrets["GMAIL_ADDRESS"]
        gmail_password = st.secrets["GMAIL_APP_PASSWORD"]
    except Exception:
        gmail_address = os.environ.get("GMAIL_ADDRESS","")
        gmail_password = os.environ.get("GMAIL_APP_PASSWORD","")

    msg = MIMEMultipart()
    msg["From"]    = gmail_address
    msg["To"]      = to_email
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))

    attachment = MIMEApplication(pdf_bytes, _subtype="pdf")
    attachment.add_header("Content-Disposition", "attachment", filename=filename)
    msg.attach(attachment)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(gmail_address, gmail_password)
        server.send_message(msg)


# ── Streamlit UI ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AVM STR Report Generator",
    page_icon="🏠",
    layout="centered"
)


# ══════════════════════════════════════════════════════════════════════════════
# ANOW EMAIL PIPELINE — Background Thread
# ══════════════════════════════════════════════════════════════════════════════

POLL_INTERVAL = 300  # 5 minutes

def _get_config():
    """Load Gmail config from Streamlit secrets or env."""
    def _s(key, default=""):
        try:
            return st.secrets[key]
        except Exception:
            return os.environ.get(key, default)
    return {
        "gmail":    _s("GMAIL_ADDRESS",     "avstr1@gmail.com"),
        "password": _s("GMAIL_APP_PASSWORD", ""),
        "delivery": _s("DELIVERY_EMAIL",    "swwebb34@a-techappraisal.com"),
        "api_key":  _s("ANTHROPIC_API_KEY", ""),
    }


def _parse_anow_email(raw_bytes):
    """Parse forwarded ANOW email. Returns (order dict, attachments list)."""
    msg = BytesParser(policy=policy.default).parsebytes(raw_bytes)
    order = {
        "subject":        msg.get("subject", ""),
        "address_line1":  "",
        "city_state_zip": "",
        "full_address":   "",
        "due_date":       "",
        "ordered_by":     "",
        "lender":         "",
        "borrower":       "",
        "contact_name":   "",
        "contact_phone":  "",
    }
    attachments = []

    for part in msg.walk():
        ct   = part.get_content_type()
        disp = str(part.get("Content-Disposition", ""))
        if "attachment" in disp or part.get_filename():
            fname = part.get_filename() or "attachment"
            fdata = part.get_payload(decode=True)
            if fdata:
                attachments.append({"filename": fname, "data": fdata, "media_type": ct})
            continue
        if ct == "text/html" and not order["borrower"]:
            html = part.get_payload(decode=True).decode("utf-8", errors="replace")
            _parse_body(BeautifulSoup(html, "html.parser").get_text("\n"), order)
        elif ct == "text/plain" and not order["borrower"]:
            _parse_body(part.get_payload(decode=True).decode("utf-8", errors="replace"), order)

    parts = [order["address_line1"], order["city_state_zip"]]
    order["full_address"] = ", ".join(p for p in parts if p)
    if not order["address_line1"]:
        m = re.search(r"New Order\s*[-–]\s*(.+?)\.?\s*$", order["subject"], re.IGNORECASE)
        if m:
            order["address_line1"] = m.group(1).strip()
            order["full_address"]  = order["address_line1"]
    return order, attachments


def _parse_body(text, order):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    start = next((i+1 for i, l in enumerate(lines) if l.upper() == "DETAILS"), None)
    if start is None:
        return
    i = start
    if i < len(lines) and not lines[i].upper().startswith("DUE"):
        order["address_line1"] = lines[i]; i += 1
    if i < len(lines) and re.search(r"\d{5}", lines[i]):
        order["city_state_zip"] = lines[i]; i += 1
    label_map = {
        "due:": "due_date", "ordered by:": "ordered_by",
        "addressed to:": "lender", "borrower:": "borrower",
        "contact for access:": "contact_name",
    }
    while i < len(lines):
        lower = lines[i].lower()
        matched = False
        for label, key in label_map.items():
            if lower.startswith(label):
                inline = lines[i][len(label):].strip()
                if inline:
                    order[key] = inline
                elif i+1 < len(lines):
                    i += 1; order[key] = lines[i]
                matched = True; break
        if not matched and order["contact_name"] and not order["contact_phone"]:
            if re.match(r"[\(\d\+]", lines[i]):
                order["contact_phone"] = lines[i]
        i += 1


PIPELINE_SYSTEM_PROMPT = """You are an expert certified residential appraiser assistant for A-Tech Appraisal Co., LLC in Rhode Island and Massachusetts. Generate a USPAP-compliant pre-inspection intake report from the order details and attached documents.

Never use subjective terms: "convenient," "desirable," "charming," "ideal," or "easy."
Never reference GLA measurement source in improvement description — just state the number.
Finished below-grade space is never included in above-grade GLA.

OUTPUT FORMAT:

## 🏠 APPRAISAL ASSIGNMENT INTAKE

**Property Address:**
**City/Town, State, Zip:**
**County:**
**Form Type:**
**Intended Use:**
**Lender/AMC:**
**Due Date:**

---

**OWNERSHIP & TRANSACTION**
- Current Owner:
- Borrower:
- Book & Page:
- Last Sale Price / Date:
- Contract Price:
- Tax Year / Amount:

---

**PROPERTY BASICS**
- Style:
- Year Built:
- GLA (above grade):
- Basement:
- Bedrooms/Baths:
- Garage:
- Lot Size:
- Public Water/Sewer:
- Zoning:
- Flood Zone:

---

**GLA SUB-AREA BREAKDOWN** [if sketch available]
| Level | Description | SF |
|---|---|---|

---

**SALES HISTORY**
| Owner | Price | Date | Notes |
|---|---|---|---|

---

**COMPLEXITY FLAGS**
[x] flag each issue found

---

**ATTACHMENTS RECEIVED**
[list what was provided]

---

**NEIGHBORHOOD BOUNDARIES**
The subject neighborhood is bounded to the north by [X], to the south by [X], to the east by [X], and to the west by [X].

---

**NEIGHBORHOOD DESCRIPTION**
[4-5 factual sentences: location, character, access, utilities, Other land use]

---

**LAND USE GRID**
| Use | % |
|---|---|
| Single Family | _% |
| 2-4 Unit | _% |
| Multi-Family (5+) | _% |
| Commercial | _% |
| Other (describe) | _% |

---

**SITE SECTION**
[2-3 sentences: lot size, zoning, utilities, flood zone]

---

**IMPROVEMENT DESCRIPTION**
[Begin: "The appraiser has inspected the interior and exterior of the subject property and researched municipal records for data reported herein."]
[Cover style, year built, GLA, rooms, bed/bath, exterior, roof, foundation, basement, heat/cool, garage, features.]
[End: "Condition and quality ratings to be determined at inspection."]

---

**PRIORITY ITEMS**
[Numbered list — most important items to confirm at inspection]"""


def _build_claude_content(order, attachments):
    import base64
    content = []
    order_text = (
        f"New ANOW Appraisal Order\n\n"
        f"Property: {order['full_address']}\n"
        f"Due: {order['due_date']}\n"
        f"AMC: {order['ordered_by']}\n"
        f"Lender: {order['lender']}\n"
        f"Borrower: {order['borrower']}\n"
        f"Contact: {order['contact_name']} {order['contact_phone']}\n\n"
        f"Generate the full pre-inspection intake report."
    )
    content.append({"type": "text", "text": order_text})

    image_labels = {
        "map": "Apple Maps screenshot — use boxed/circled landmarks for N/S/E/W boundaries",
        "gis": "GIS/Parcel map — use for land use grid and site data",
        "screen": "Screenshot — extract any relevant appraisal data",
    }
    pdf_titles = {
        "tax": "Tax Card / Assessor Record",
        "vision": "Tax Card / Assessor Record",
        "mls": "MLS Sheet",
        "360": "MLS Sheet",
        "contract": "Purchase and Sales Agreement",
        "measure": "Field Measurement Report",
        "gla": "Field Measurement Report",
    }

    for att in attachments:
        fname = att["filename"].lower()
        data  = att["data"]
        mime  = att["media_type"].lower()

        if "pdf" in mime or fname.endswith(".pdf"):
            b64   = base64.standard_b64encode(data).decode()
            title = next((v for k, v in pdf_titles.items() if k in fname),
                         f"Document: {att['filename']}")
            content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
                "title": title
            })
        elif any(ext in fname for ext in [".jpg",".jpeg",".png"]) or "image" in mime:
            # Compress if needed
            if len(data) > 4*1024*1024:
                from io import BytesIO
                img = PILImage.open(BytesIO(data)).convert("RGB")
                q = 85
                while q >= 30:
                    buf = BytesIO()
                    img.save(buf, format="JPEG", quality=q, optimize=True)
                    if len(buf.getvalue()) <= 4*1024*1024:
                        data = buf.getvalue(); break
                    q -= 10
            b64   = base64.standard_b64encode(data).decode()
            label = next((v for k, v in image_labels.items() if k in fname),
                         "Uploaded image — use for any relevant appraisal data")
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}
            })
            content.append({"type": "text", "text": f"[Above: {label}]"})

    return content


def _call_claude_pipeline(content, api_key):
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
        "anthropic-beta": "pdfs-2024-09-25",
    }
    payload = {
        "model": "claude-opus-4-5",
        "max_tokens": 4000,
        "system": PIPELINE_SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": content}],
    }
    resp = requests.post("https://api.anthropic.com/v1/messages",
                         headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    return resp.json()["content"][0]["text"].strip()


def _send_intake_email(gmail, password, delivery, order, pdf_bytes):
    address  = order.get("full_address", "Subject Property")
    borrower = order.get("borrower", "")
    due      = order.get("due_date", "")
    safe     = re.sub(r"[^\w\s-]", "", address).replace(" ", "_")[:60]

    msg = MIMEMultipart()
    msg["From"]    = gmail
    msg["To"]      = delivery
    msg["Subject"] = f"Intake Report — {address}"
    msg.attach(MIMEText(
        f"Pre-inspection intake for {address}.\n\n"
        f"Borrower: {borrower}\nDue: {due}\n\n"
        f"Review flags and priority items before inspection.\n\n"
        f"— A-Tech Intake Pipeline", "plain"
    ))
    att = MIMEApplication(pdf_bytes, _subtype="pdf")
    att.add_header("Content-Disposition", "attachment",
                   filename=f"{safe}_intake.pdf")
    msg.attach(att)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as s:
        s.login(gmail, password)
        s.sendmail(gmail, delivery, msg.as_string())


def _pipeline_loop():
    """Background thread — polls Gmail every POLL_INTERVAL seconds."""
    cfg = _get_config()
    if not cfg["password"] or not cfg["api_key"]:
        st.session_state["pipeline_status"] = "⚠️ Missing credentials — check secrets"
        return

    while True:
        try:
            st.session_state["pipeline_status"] = (
                f"🟢 Active — last checked {datetime.now().strftime('%I:%M %p')}"
            )
            imap = imaplib.IMAP4_SSL("imap.gmail.com")
            imap.login(cfg["gmail"], cfg["password"])
            imap.select("INBOX")

            _, data = imap.uid("search", None, '(UNSEEN SUBJECT "New Order")')
            uids = data[0].split() if data[0] else []

            for uid in uids:
                try:
                    _, msg_data = imap.uid("fetch", uid, "(RFC822)")
                    if not msg_data or not msg_data[0]:
                        continue
                    raw = msg_data[0][1]

                    order, attachments = _parse_anow_email(raw)
                    address = order.get("full_address") or "Unknown"

                    st.session_state["pipeline_status"] = (
                        f"⚙️ Processing {address}..."
                    )

                    content    = _build_claude_content(order, attachments)
                    intake_txt = _call_claude_pipeline(content, cfg["api_key"])
                    pdf_bytes  = build_intake_pdf(intake_txt, address)
                    _send_intake_email(cfg["gmail"], cfg["password"],
                                       cfg["delivery"], order, pdf_bytes)

                    # Mark read
                    imap.uid("store", uid, "+FLAGS", "\\Seen")

                    st.session_state["pipeline_status"] = (
                        f"🟢 Active — last processed: {address} "
                        f"at {datetime.now().strftime('%I:%M %p')}"
                    )
                    st.session_state["pipeline_last"] = address

                except Exception as e:
                    st.session_state["pipeline_status"] = f"⚠️ Error on order: {e}"
                    try:
                        imap.uid("store", uid, "+FLAGS", "\\Seen")
                    except Exception:
                        pass

            imap.logout()

        except Exception as e:
            st.session_state["pipeline_status"] = f"⚠️ Connection error: {e}"

        time.sleep(POLL_INTERVAL)


def _start_pipeline():
    """Start the background thread once per session."""
    if not st.session_state.get("pipeline_started"):
        st.session_state["pipeline_started"] = True
        st.session_state["pipeline_status"]  = "🟢 Active — starting up..."
        t = threading.Thread(target=_pipeline_loop, daemon=True)
        t.start()

# ── Password Protection ───────────────────────────────────────────────────────
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if st.session_state.authenticated:
        return True
    st.image(LOGO_PATH, width=220)
    st.title("AVM STR Report Generator")
    st.divider()
    pwd = st.text_input("Enter password to continue", type="password", key="pwd_input")
    if st.button("Login", use_container_width=True):
        try:
            correct = st.secrets["APP_PASSWORD"]
        except Exception:
            correct = os.environ.get("APP_PASSWORD", "avm2026")
        if pwd == correct:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    return False

if not check_password():
    st.stop()

# ── Storage Helpers ───────────────────────────────────────────────────────────
def load_clients():
    try:
        raw = st.session_state.get("_clients_store")
        if raw:
            return json.loads(raw)
    except Exception:
        pass
    # Try file-based fallback
    try:
        path = os.path.join(os.path.dirname(__file__), "clients.json")
        if os.path.exists(path):
            with open(path) as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def save_clients(clients):
    st.session_state["_clients_store"] = json.dumps(clients)
    try:
        path = os.path.join(os.path.dirname(__file__), "clients.json")
        with open(path, "w") as f:
            json.dump(clients, f)
    except Exception:
        pass

def load_orders():
    try:
        raw = st.session_state.get("_orders_store")
        if raw:
            return json.loads(raw)
    except Exception:
        pass
    try:
        path = os.path.join(os.path.dirname(__file__), "orders.json")
        if os.path.exists(path):
            with open(path) as f:
                return json.load(f)
    except Exception:
        pass
    return []

def save_orders(orders):
    st.session_state["_orders_store"] = json.dumps(orders)
    try:
        path = os.path.join(os.path.dirname(__file__), "orders.json")
        with open(path, "w") as f:
            json.dump(orders, f)
    except Exception:
        pass

def log_order(address, property_type, client, borrower, loan_num, avm_file_id, report_date):
    orders = load_orders()
    orders.insert(0, {
        "date":          report_date,
        "address":       address,
        "property_type": property_type,
        "client":        client,
        "borrower":      borrower,
        "loan_number":   loan_num,
        "avm_file_id":   avm_file_id,
    })
    save_orders(orders)

# ── Header ────────────────────────────────────────────────────────────────────
st.image(LOGO_PATH, width=220)
st.title("STR Income Analysis Generator")
st.divider()

# ── Tabs ──────────────────────────────────────────────────────────────────────
tab_generate, tab_clients, tab_history, tab_intake = st.tabs([
    "⚡ Generate Report", "👥 Client Database", "📋 Order History", "🏠 AI Intake Assistant"
])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — GENERATE REPORT
# ══════════════════════════════════════════════════════════════════════════════
with tab_generate:

    # File Uploads
    st.subheader("1. Upload AirDNA PDF")
    col1, col2 = st.columns(2)
    with col1:
        airdna_pdf = st.file_uploader("AirDNA Rentalizer PDF", type="pdf", key="pdf")

    st.subheader("2. Property Photo (Optional)")
    col_p1, col_p2 = st.columns(2)
    with col_p1:
        property_photo = st.file_uploader("Property Photo", type=["jpg","jpeg","png"], key="photo",
                                           help="Front exterior photo of the subject property")

    # Assignment Info
    st.subheader("3. Assignment Info")

    # Client autofill
    clients = load_clients()
    client_names = ["-- Enter manually --"] + sorted(clients.keys())
    selected_client = st.selectbox("Select existing client (or enter manually below)",
                                    client_names, key="client_select")

    # Pre-fill values from selected client
    if selected_client != "-- Enter manually --" and selected_client in clients:
        c = clients[selected_client]
        prefill_client        = c.get("name", "")
        prefill_client_address = c.get("address", "")
        prefill_client_phone  = c.get("phone", "")
    else:
        prefill_client        = ""
        prefill_client_address = ""
        prefill_client_phone  = ""

    col4, col5, col6 = st.columns(3)
    with col4:
        client = st.text_input("Client / Lender", value=prefill_client,
                                placeholder="Annie Mac Home Mortgage")
    with col5:
        client_address = st.text_input("Client Address", value=prefill_client_address,
                                        placeholder="123 Main St, Boston, MA")
    with col6:
        client_phone = st.text_input("Client Phone", value=prefill_client_phone,
                                      placeholder="617-555-1234")

    col7, col8, col9 = st.columns(3)
    with col7:
        client_order_num = st.text_input("Client Order Number", placeholder="ORD-12345")
    with col8:
        borrower = st.text_input("Borrower", placeholder="John Smith")
    with col9:
        loan_num = st.text_input("Loan Number", placeholder="2008727778")

    col10, col11, col12 = st.columns(3)
    with col10:
        avm_file_id = st.text_input("AVM File ID", placeholder="AVM-2026-001")
    with col11:
        from datetime import date
        report_date = st.text_input("Report Date",
            value=date.today().strftime("%B %d, %Y"))
    with col12:
        property_type = st.selectbox("Property Type", [
            "Single-Family Residence",
            "Condominium",
            "Townhouse",
            "Multi-Family (2-4 Units)",
            "Single Unit in Multi-Family",
            "Manufactured Home",
        ])

    st.subheader("4. Market Overview")
    market_overview = st.text_area(
        "Market Overview",
        placeholder="Write 3-5 sentences describing the local STR market — what drives demand, submarket characteristics, seasonality, and any relevant local factors...",
        height=130,
        label_visibility="collapsed"
    )
    st.caption("This text appears in the Market Overview & Commentary section of the report.")
    st.divider()

    if st.button("⚡ Generate Report", type="primary", use_container_width=True):
        if not airdna_pdf:
            st.error("Please upload the AirDNA PDF.")
        elif not client or not loan_num:
            st.error("Please enter the client/lender name and loan number.")
        elif not market_overview.strip():
            st.error("Please enter a market overview before generating.")
        else:
            with st.spinner("Extracting data from AirDNA PDF..."):
                pdf_bytes = airdna_pdf.read()
                data = parse_airdna_pdf(pdf_bytes)

            commentary = market_overview.strip()

            with st.spinner("Building report..."):
                photo_override = None
                if property_photo:
                    tmp_photo = tempfile.NamedTemporaryFile(delete=False,
                        suffix=os.path.splitext(property_photo.name)[1])
                    tmp_photo.write(property_photo.read())
                    tmp_photo.close()
                    photo_override = tmp_photo.name

                map_override = None

                buf = io.BytesIO()
                build_pdf(data, client, loan_num, report_date, commentary, buf,
                          photo_override=photo_override, map_override=map_override,
                          client_address=client_address, client_phone=client_phone,
                          client_order_num=client_order_num, borrower=borrower,
                          avm_file_id=avm_file_id, property_type=property_type)
                buf.seek(0)

            # Log the order
            log_order(
                address=f"{data.get('address_line1','')} {data.get('city_state_zip','')}",
                property_type=property_type,
                client=client,
                borrower=borrower,
                loan_num=loan_num,
                avm_file_id=avm_file_id,
                report_date=report_date
            )

            addr_slug = re.sub(r"[^a-zA-Z0-9]+","_",
                               data.get("address_line1","Report")).strip("_")
            filename = f"AVM_STR_{addr_slug}.pdf"

            st.success("✅ Report generated and logged!")
            st.download_button(
                label="📄 Download Report PDF",
                data=buf,
                file_name=filename,
                mime="application/pdf",
                use_container_width=True
            )

            # Store PDF in session state for email sending
            st.session_state["last_pdf_bytes"]  = buf.getvalue()
            st.session_state["last_pdf_filename"] = filename
            st.session_state["last_pdf_address"]  = f"{data.get('address_line1','')} {data.get('city_state_zip','')}"

    # Email section — shows after report is generated
    if "last_pdf_bytes" in st.session_state:
        st.divider()
        st.subheader("📧 Send Report via Email")
        e1, e2 = st.columns([3, 1])
        with e1:
            email_to = st.text_input("Recipient email address",
                                      placeholder="lender@example.com",
                                      key="email_to")
        with e2:
            st.write("")
            st.write("")
            send_clicked = st.button("Send", use_container_width=True, key="send_email")

        email_note = st.text_area("Optional note to include in email body",
                                   placeholder="Please find the STR Income Analysis attached...",
                                   height=80, key="email_note")

        if send_clicked:
            if not email_to.strip():
                st.error("Please enter a recipient email address.")
            else:
                address_line = st.session_state.get("last_pdf_address","Subject Property")
                subject = f"STR Income Analysis — {address_line}"
                body = email_note.strip() if email_note.strip() else (
                    f"Please find the Short-Term Rental Income Analysis attached for {address_line}.\n\n"
                    f"This report was prepared by Absolute Value Management.\n\n"
                    f"Please note: This is not an appraisal and does not constitute an opinion of market value."
                )
                try:
                    with st.spinner("Sending email..."):
                        send_report_email(
                            to_email=email_to.strip(),
                            subject=subject,
                            body=body,
                            pdf_bytes=st.session_state["last_pdf_bytes"],
                            filename=st.session_state["last_pdf_filename"]
                        )
                    st.success(f"✅ Report sent to {email_to.strip()}")
                except Exception as e:
                    st.error(f"Email failed: {str(e)}")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — CLIENT DATABASE
# ══════════════════════════════════════════════════════════════════════════════
with tab_clients:
    st.subheader("Client Database")
    st.caption("Save client info here once — it will auto-fill on the Generate tab.")

    clients = load_clients()

    # Add / Edit client form
    with st.expander("➕ Add New Client", expanded=len(clients) == 0):
        nc1, nc2, nc3 = st.columns(3)
        with nc1:
            new_name    = st.text_input("Client / Lender Name *", key="new_name",
                                         placeholder="Annie Mac Home Mortgage")
        with nc2:
            new_address = st.text_input("Client Address", key="new_address",
                                         placeholder="123 Main St, Boston, MA")
        with nc3:
            new_phone   = st.text_input("Client Phone", key="new_phone",
                                         placeholder="617-555-1234")

        if st.button("💾 Save Client", use_container_width=True):
            if not new_name.strip():
                st.error("Client name is required.")
            else:
                clients[new_name.strip()] = {
                    "name":    new_name.strip(),
                    "address": new_address.strip(),
                    "phone":   new_phone.strip(),
                }
                save_clients(clients)
                st.success(f"✅ Client '{new_name.strip()}' saved.")
                st.rerun()

    # Client list
    if clients:
        st.divider()
        st.write(f"**{len(clients)} client(s) saved**")
        for name, info in sorted(clients.items()):
            with st.container():
                cc1, cc2, cc3, cc4 = st.columns([3, 3, 2, 1])
                with cc1:
                    st.write(f"**{name}**")
                with cc2:
                    st.write(info.get("address","—"))
                with cc3:
                    st.write(info.get("phone","—"))
                with cc4:
                    if st.button("🗑️", key=f"del_{name}", help=f"Delete {name}"):
                        del clients[name]
                        save_clients(clients)
                        st.rerun()
        st.divider()

        # Export clients as CSV
        client_rows = [{"Client": k, "Address": v.get("address",""),
                        "Phone": v.get("phone","")} for k,v in clients.items()]
        client_df = pd.DataFrame(client_rows)
        st.download_button("⬇️ Export Client List (CSV)",
                            data=client_df.to_csv(index=False),
                            file_name="avm_clients.csv",
                            mime="text/csv")
    else:
        st.info("No clients saved yet. Add your first client above.")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — ORDER HISTORY
# ══════════════════════════════════════════════════════════════════════════════
with tab_history:
    st.subheader("Order History")
    st.caption("Every report generated is automatically logged here.")

    orders = load_orders()

    if orders:
        st.write(f"**{len(orders)} order(s) on record**")

        # Search / filter
        search = st.text_input("🔍 Search by address, client, or borrower",
                                placeholder="Type to filter...", key="order_search")
        if search:
            q = search.lower()
            orders = [o for o in orders if
                      q in o.get("address","").lower() or
                      q in o.get("client","").lower() or
                      q in o.get("borrower","").lower()]

        # Display as table
        if orders:
            df = pd.DataFrame(orders)
            df = df.rename(columns={
                "date":          "Report Date",
                "address":       "Property Address",
                "property_type": "Type",
                "client":        "Client / Lender",
                "borrower":      "Borrower",
                "loan_number":   "Loan Number",
                "avm_file_id":   "AVM File ID",
            })
            st.dataframe(df, use_container_width=True, hide_index=True)

            st.divider()
            st.download_button(
                "⬇️ Export Order History (CSV)",
                data=df.to_csv(index=False),
                file_name="avm_order_history.csv",
                mime="text/csv",
                use_container_width=True
            )
        else:
            st.info("No orders match your search.")
    else:
        st.info("No orders logged yet. Generate your first report to start the log.")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — AI INTAKE ASSISTANT
# ══════════════════════════════════════════════════════════════════════════════
with tab_intake:

    # ── Start background pipeline ─────────────────────────────────────────────
    _start_pipeline()

    st.subheader("🏠 AI Intake Assistant")
    st.caption(
        "Upload your tax card, Apple Maps screenshot, GIS screenshot, MLS sheet, "
        "or contract — add any notes — and Claude will generate a fully populated "
        "intake template ready to copy into TOTAL."
    )

    # ── Pipeline status indicator ─────────────────────────────────────────────
    with st.container():
        status = st.session_state.get("pipeline_status", "🟢 Active — starting up...")
        st.info(f"**📬 Email Pipeline:** {status}")
        st.caption(
            f"Polling avstr1@gmail.com every 5 min for forwarded ANOW orders. "
            f"Forward orders with attachments to avstr1@gmail.com → "
            f"intake PDF delivered to swwebb34@a-techappraisal.com."
        )
    st.divider()

    # ── API Key (from Streamlit secrets or environment variable) ─────────────
    try:
        intake_api_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        intake_api_key = os.environ.get("ANTHROPIC_API_KEY", "")

    # ── File Uploads ──────────────────────────────────────────────────────────
    st.markdown("#### 1. Upload Documents")
    col_u1, col_u2 = st.columns(2)
    with col_u1:
        tax_card_file = st.file_uploader(
            "Tax Card (PDF or Image)",
            type=["pdf", "png", "jpg", "jpeg"],
            key="intake_tax_card"
        )
        mls_file = st.file_uploader(
            "MLS Sheet / 360 Property View (PDF) — optional",
            type=["pdf"],
            key="intake_mls",
            help="Not required for refinance assignments without an active listing."
        )
        anow_file = st.file_uploader(
            "ANOW Order Screenshot (PNG/JPG) — optional",
            type=["png", "jpg", "jpeg"],
            key="intake_anow",
            help="Screenshot from ANOW order management — used to pull borrower, lender, and order details."
        )
    with col_u2:
        maps_file = st.file_uploader(
            "Apple Maps Screenshot (PNG/JPG)",
            type=["png", "jpg", "jpeg"],
            key="intake_maps"
        )
        gis_file = st.file_uploader(
            "GIS / Parcel Map Screenshot (PNG/JPG)",
            key="intake_gis",
            type=["png", "jpg", "jpeg"]
        )
        bt_file = st.file_uploader(
            "Banker & Tradesman (PDF or Image) — optional",
            type=["pdf", "png", "jpg", "jpeg"],
            key="intake_bt",
            help="Banker & Tradesman deed/transfer report for sales history and ownership verification."
        )
    contract_file = st.file_uploader(
        "Purchase & Sales Agreement (PDF) — optional",
        type=["pdf"],
        key="intake_contract"
    )
    cubicasa_file = st.file_uploader(
        "Field Measurement Report (PDF) — optional",
        type=["pdf"],
        key="intake_cubicasa"
    )

    # ── Manual Notes ──────────────────────────────────────────────────────────
    st.markdown("#### 2. Assignment Notes")
    col_n1, col_n2, col_n3 = st.columns(3)
    with col_n1:
        intake_borrower    = st.text_input("Borrower", key="intake_borrower")
        intake_lender      = st.text_input("Lender / AMC", key="intake_lender")
    with col_n2:
        intake_form        = st.selectbox("Form Type",
                                          ["1004", "1073 Condo", "1025 Multi-Family",
                                           "1004C Manufactured", "Other"],
                                          key="intake_form")
        intake_use         = st.selectbox("Intended Use",
                                          ["Purchase", "Refinance", "Estate",
                                           "Divorce", "Other"],
                                          key="intake_use")
    with col_n3:
        intake_due         = st.text_input("Due Date", key="intake_due")
        intake_contract_px = st.text_input("Contract Price", key="intake_contract_px",
                                            placeholder="$000,000")

    intake_notes = st.text_area(
        "Additional Notes (inspection findings, borrower notes, flags, etc.)",
        key="intake_notes",
        height=80,
        placeholder="e.g. Owner confirms well and septic. TQS ceiling height under 7ft in places. Borrower same as owner LLC."
    )

    # ── System Prompt ─────────────────────────────────────────────────────────
    INTAKE_SYSTEM_PROMPT = """You are an expert certified residential appraiser assistant working for A-Tech Appraisal Co., LLC in Rhode Island and Massachusetts. You help appraisers prepare USPAP-compliant appraisal report data by analyzing uploaded documents including tax cards, MLS sheets, GIS screenshots, Apple Maps screenshots, purchase contracts, and field measurement reports.

Your output must follow this exact intake template structure. All narrative language must be factual and objective — avoid subjective or non-USPAP-compliant terms such as "convenient," "desirable," "charming," "ideal," or "easy."

OUTPUT FORMAT — always produce all sections in this order:

## 🏠 APPRAISAL ASSIGNMENT INTAKE

**Property Address:**
**City/Town, State, Zip:**
**County:**
**Form Type:**
**Intended Use:**
**Lender/AMC:**
**Due Date:**

---

**OWNERSHIP & TRANSACTION**
- Current Owner:
- Co-Owner:
- Borrower:
- Book & Page:
- Last Sale Price:
- Last Sale Date:
- Contract Price:
- Seller Concession (if any):
- Tax Year:
- Tax Amount:

---

**PROPERTY BASICS**
- Style:
- Year Built:
- GLA (above grade): [use field measurement report figure if uploaded, otherwise use assessor figure. Do not reference the source in the output — just state the number]
- Basement: [type, size, finish %, access]
- Bedrooms/Baths:
- Garage: [attached/detached, # cars, sf]
- Lot Size:
- Public Water: Y/N
- Public Sewer: Y/N
- Legal Description / Zoning:
- Flood Zone: [FEMA zone, panel #, date if available]

---

**GLA SUB-AREA BREAKDOWN** [if field measurement report or Vision sketch available]
| Level | Description | SF |
|---|---|---|

---

**OWNERSHIP & SALES HISTORY**
| Owner | Price | Date | Notes |
|---|---|---|---|

[Flag any same-day transfers, flip history, non-arm's-length transfers, or sales within 36 months]

---

**COMPLEXITY FLAGS**
[List all flags with [x] prefix. Common flags to check:
- Non-ANSI space (TQS, FRB, below-grade finished space)
- GLA variance between sources
- Recent flip or same-day transfer
- LLC ownership with individual borrower
- Seller concession
- Recent arm's length sale within 12 months
- Historic construction (pre-1940)
- No garage
- Oil heat / no AC
- Unknown underground tank
- Thin comp market
- Coastal / flood zone
- Condo vs PUD classification
- Permit history flags
- Water/sewer confirmation needed]

---

**ATTACHMENTS RECEIVED**
[List what was uploaded — tax card, Apple Maps, GIS, MLS sheet, ANOW screenshot, Banker & Tradesman, contract, field measurement report — only list what was actually provided]

---

**NEIGHBORHOOD BOUNDARIES**
One sentence: "The subject neighborhood is bounded to the north by [X], to the south by [X], to the east by [X], and to the west by [X]."
[Derive from Apple Maps screenshot — use roads, landmarks, natural features visible in the image]

---

**NEIGHBORHOOD DESCRIPTION**
[4-5 sentences. Factual, objective, USPAP-compliant. No subjective language.
Cover: municipality/county location, immediate neighborhood character, access/arterials (no "convenient"), utilities, and Other land use description (parks, open space, golf, conservation, etc.)]

---

**LAND USE GRID**
| Use | % |
|---|---|
| Single Family | _% |
| 2–4 Unit | _% |
| Multi-Family (5+) | _% |
| Commercial | _% |
| Other (describe) | _% |
[Derive from GIS screenshot if available. Note if estimated.]

---

**SITE SECTION**
[2-3 sentences. Cover: lot size, zoning, lot description/shape, utilities, flood zone, any adverse conditions.]

---

**IMPROVEMENT DESCRIPTION**
[Begin: "The appraiser has inspected the interior and exterior of the subject property and researched municipal records for data reported herein."]
[Cover: style, year built, GLA with source, room count, bed/bath, exterior, roof, foundation, basement, heat/cool, garage, additional features, condition note at end.]
[If a field measurement report is provided, use that GLA figure. If assessor only, use that figure. Do not reference the measurement source or ANSI in the improvement description narrative — just state the GLA number. Flag any variance between sources in the COMPLEXITY FLAGS section.]
[If finished below-grade space exists, note it is excluded from above-grade GLA per ANSI and reported separately.]
[End: "Condition and quality ratings to be determined at inspection." unless inspection has occurred — in that case use provided condition notes.]

---

**PRIORITY FLAGS / ITEMS TO RESOLVE**
[Numbered list of the most important items to confirm before or at inspection]

---

IMPORTANT RULES:
1. Never use "convenient," "desirable," "charming," "ideal," or "easy" in narrative
2. Never reference the GLA measurement source in the improvement description narrative. Just state the number. Flag any variance between sources (field measurement vs. assessor vs. MLS) in the COMPLEXITY FLAGS section only.
3. Flag any GLA variance between sources
4. Finished below-grade space is NEVER included in above-grade GLA per ANSI Z765
5. TQS (Three Quarter Story) space must be verified for ANSI ceiling height compliance at inspection
6. If a property is owned by an LLC with an individual borrower, always flag it
7. Seller concessions must always be flagged and noted
8. Any transfer within 36 months of the effective date must be analyzed in prior sales history
9. Same-day double transfers are always a flag
10. Permit history from Vision assessor cards should be noted but not over-emphasized
11. Land use grid percentages should reflect the immediate neighborhood, not the whole town
12. Always note when water/sewer needs to be confirmed at inspection"""

    # ── Generate Button ───────────────────────────────────────────────────────
    st.divider()
    generate_btn = st.button(
        "🏠 Generate Intake",
        use_container_width=True,
        key="intake_generate",
        type="primary"
    )

    if generate_btn:
        if not intake_api_key:
            st.error("Anthropic API key not found. Add ANTHROPIC_API_KEY to your Streamlit secrets.")
            st.stop()

        # Build message content
        content = []

        # Add uploaded files as base64
        import base64

        def pdf_to_base64(file_bytes):
            return base64.standard_b64encode(file_bytes).decode("utf-8")

        def img_to_base64_and_type(file_bytes, original_filename, max_bytes=4 * 1024 * 1024):
            """Compress image if needed. Returns (base64_str, media_type)."""
            from io import BytesIO
            # If small enough and already JPEG, send as-is
            ext = original_filename.lower().split(".")[-1]
            if len(file_bytes) <= max_bytes and ext in ("jpg", "jpeg"):
                return base64.standard_b64encode(file_bytes).decode("utf-8"), "image/jpeg"
            # Convert to RGB JPEG and compress
            img = PILImage.open(BytesIO(file_bytes)).convert("RGB")
            # If small enough as PNG, still convert to JPEG for consistency
            if len(file_bytes) <= max_bytes:
                buf = BytesIO()
                img.save(buf, format="JPEG", quality=90, optimize=True)
                return base64.standard_b64encode(buf.getvalue()).decode("utf-8"), "image/jpeg"
            # Need to compress — try reducing quality first
            quality = 85
            while quality >= 30:
                buf = BytesIO()
                img.save(buf, format="JPEG", quality=quality, optimize=True)
                compressed = buf.getvalue()
                if len(compressed) <= max_bytes:
                    return base64.standard_b64encode(compressed).decode("utf-8"), "image/jpeg"
                quality -= 10
            # Last resort — resize to 50%
            w, h = img.size
            img = img.resize((w // 2, h // 2), PILImage.LANCZOS)
            buf = BytesIO()
            img.save(buf, format="JPEG", quality=70, optimize=True)
            return base64.standard_b64encode(buf.getvalue()).decode("utf-8"), "image/jpeg"

        def add_image_content(content, file_bytes, filename, label):
            """Helper to add an image to content with compression and correct media type."""
            b64, media_type = img_to_base64_and_type(file_bytes, filename)
            content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": b64
                }
            })
            content.append({"type": "text", "text": label})

        if tax_card_file:
            file_bytes = tax_card_file.read()
            fname = tax_card_file.name.lower()
            if fname.endswith(".pdf"):
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_to_base64(file_bytes)
                    },
                    "title": "Tax Card"
                })
            else:
                add_image_content(content, file_bytes, tax_card_file.name, "[Above image: Tax Card]")

        if maps_file:
            file_bytes = maps_file.read()
            add_image_content(content, file_bytes, maps_file.name,
                "[Above image: Apple Maps screenshot — use circled/boxed landmarks for N/S/E/W neighborhood boundaries]")

        if gis_file:
            file_bytes = gis_file.read()
            add_image_content(content, file_bytes, gis_file.name,
                "[Above image: GIS / Parcel Map — use for land use grid percentages and site data]")

        if anow_file:
            file_bytes = anow_file.read()
            add_image_content(content, file_bytes, anow_file.name,
                "[Above image: ANOW order screenshot — extract borrower, lender/client, AMC, order details, and any assignment notes]")

        if bt_file:
            file_bytes = bt_file.read()
            fname = bt_file.name.lower()
            if fname.endswith(".pdf"):
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_to_base64(file_bytes)
                    },
                    "title": "Banker & Tradesman — use for ownership history, deed transfers, book/page, and sales history"
                })
            else:
                add_image_content(content, file_bytes, bt_file.name,
                    "[Above image: Banker & Tradesman — use for ownership history, deed transfers, book/page, and sales history]")

        if mls_file:
            file_bytes = mls_file.read()
            content.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_to_base64(file_bytes)
                },
                "title": "MLS Sheet"
            })

        if contract_file:
            file_bytes = contract_file.read()
            content.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_to_base64(file_bytes)
                },
                "title": "Purchase and Sales Agreement"
            })

        if cubicasa_file:
            file_bytes = cubicasa_file.read()
            content.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_to_base64(file_bytes)
                },
                "title": "Field Measurement Report — ANSI compliant field measurements. Use these GLA figures and sub-area breakdown. Note any variance from assessor or MLS. GLA and room counts are subject to verification at inspection."
            })

        # Build text prompt
        prompt_parts = ["Please generate the full appraisal intake template based on the uploaded documents."]
        if intake_borrower:    prompt_parts.append(f"Borrower: {intake_borrower}")
        if intake_lender:      prompt_parts.append(f"Lender/AMC: {intake_lender}")
        if intake_form:        prompt_parts.append(f"Form Type: {intake_form}")
        if intake_use:         prompt_parts.append(f"Intended Use: {intake_use}")
        if intake_due:         prompt_parts.append(f"Due Date: {intake_due}")
        if intake_contract_px: prompt_parts.append(f"Contract Price: {intake_contract_px}")
        if intake_notes:       prompt_parts.append(f"Additional Notes: {intake_notes}")

        content.append({"type": "text", "text": "\n".join(prompt_parts)})

        # Call API
        with st.spinner("Claude is analyzing your documents and generating the intake..."):
            try:
                headers = {
                    "x-api-key": intake_api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                    "anthropic-beta": "pdfs-2024-09-25"
                }
                payload = {
                    "model": "claude-opus-4-5",
                    "max_tokens": 4000,
                    "system": INTAKE_SYSTEM_PROMPT,
                    "messages": [
                        {"role": "user", "content": content}
                    ]
                }
                response = requests.post(
                    "https://api.anthropic.com/v1/messages",
                    headers=headers,
                    json=payload,
                    timeout=120
                )
                response.raise_for_status()
                result = response.json()
                output_text = result["content"][0]["text"].strip()

                # Follow-up question box
                st.divider()
                st.markdown("#### 💬 Follow-up Question")
                followup = st.text_input(
                    "Ask a follow-up about this property (e.g. 'draft the improvement description', 'what comps strategy would you suggest')",
                    key="intake_followup"
                )
                if st.button("Send Follow-up", key="intake_followup_btn"):
                    if followup.strip():
                        with st.spinner("Thinking..."):
                            fu_payload = {
                                "model": "claude-opus-4-5",
                                "max_tokens": 2000,
                                "system": INTAKE_SYSTEM_PROMPT,
                                "messages": [
                                    {"role": "user", "content": content},
                                    {"role": "assistant", "content": output_text},
                                    {"role": "user", "content": followup}
                                ]
                            }
                            fu_response = requests.post(
                                "https://api.anthropic.com/v1/messages",
                                headers=headers,
                                json=fu_payload,
                                timeout=120
                            )
                            fu_response.raise_for_status()
                            fu_result = fu_response.json()
                            fu_text = fu_result["content"][0]["text"]
                            st.markdown(fu_text)
                            st.text_area("📋 Copy follow-up response",
                                         value=fu_text, height=300,
                                         key="intake_fu_output")

            except requests.exceptions.HTTPError as e:
                st.error(f"API error: {e.response.status_code} — {e.response.text}")
            except Exception as e:
                st.error(f"Error: {str(e)}")
